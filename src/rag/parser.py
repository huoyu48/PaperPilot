"""Document parser: extracts text from PDF, TXT, MD, DOCX files."""

from __future__ import annotations

from pathlib import Path

from langchain_core.documents import Document

from src.utils.logging import logger


def parse_document(file_path: str | Path) -> Document:
    """Parse a file into a LangChain Document.

    Supports: .pdf, .txt, .md, .docx
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    suffix = path.suffix.lower()
    logger.info(f"Parser: loading {path.name} ({suffix})")

    if suffix == ".txt" or suffix == ".md":
        text = path.read_text(encoding="utf-8", errors="replace")
    elif suffix == ".pdf":
        text = _parse_pdf(path)
    elif suffix == ".docx":
        text = _parse_docx(path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")

    return Document(
        page_content=text,
        metadata={"source": str(path), "filename": path.name, "type": suffix},
    )


def _parse_pdf(path: Path) -> str:
    """Extract text from PDF using PyPDF2 or pdfplumber."""
    try:
        import pdfplumber
        texts = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    texts.append(t)
        return "\n\n".join(texts)
    except ImportError:
        pass

    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)
    except ImportError:
        raise ImportError("Install pdfplumber or pypdf for PDF support: pip install pdfplumber")


def _parse_docx(path: Path) -> str:
    """Extract text from DOCX."""
    try:
        import docx
        doc = docx.Document(str(path))
        return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except ImportError:
        raise ImportError("Install python-docx for DOCX support: pip install python-docx")
